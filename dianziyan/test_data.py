#！C:\Users\alp-sof\AppData\Local\Programs\Python\Python39\python.exe
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import serial
import os
import sys
import numpy as np

# 声明一个word对象
doc = Document()
# 设置字体样式
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
 # ------添加文档标题-------
paragraph = doc.add_paragraph()
run = paragraph.add_run("Test Result")
font = run.font
 # 设置字体大小
font.size = Pt(15)
 # 设置水平居中
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

table = doc.add_table(rows=51, cols=9, style="Table Grid")
# table.cell(0,0).merge(table.cell(0,5))   #合并单元格
# table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.cell(0,0).text ="Item"
table.cell(0,1).text ="静态电流（uA）VDD=4V"
table.cell(0,2).text ="欠压保护（V）"
table.cell(0,3).text ="恒压输出（V）"
table.cell(0,4).text ="超时保护时间(S)"
table.cell(0,5).text ="最小导通占空比VDD=4.2V"
table.cell(0,6).text ="短路保护"
table.cell(0,7).text ="SW变化到AT输出延时（ms空载）"
table.cell(0,8).text ="芯片内阻RL=0.8Ω（mΩ）"
for i in range (1,51):
    table.cell(i,0).text=str(i)
#table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER    # 第一行表格水平居中
doc.add_paragraph("")
doc.add_paragraph("LED灯状态指示对比")
table2 = doc.add_table(rows=6, cols=3, style="Table Grid")
table2.cell(0,0).text ="电子烟状态"
table2.cell(1,0).text ="上电"
table2.cell(2,0).text ="正常吸烟"
table2.cell(3,0).text ="长时间吸烟(10秒)保护"
table2.cell(4,0).text ="过流保护"
table2.cell(5,0).text ="欠压保护"
table2.cell(0,1).text ="LED显示"

com = serial.Serial("com12", 115200)

if not com.is_open:
  print("com open failed")
  exit()

print("等待串口")
for cnt in range(1,51):
    #doc.save('E:/E/PM/测试报告/test result report.docx')
    for i in range(0, 30):
        string = com.readline()
        str1 = string.decode('utf-8')
        print(i,": ", str1)
        if  '最大占空比' in str1:
            duty=str1[str1.rfind('：')+1:-2]
            table.cell(cnt,5).text =duty
        if '恒压输出' in str1:
            cv=str1[str1.rfind('：')+1:-2]
            table.cell(cnt,3).text =cv
        if '欠压阈值' in str1:
            uvlo=str1[str1.rfind('：')+1:-2]  
            table.cell(cnt,2).text =uvlo
        if '超时保护时间' in str1:
            ot=str1[str1.rfind('：')+1:-2]
            table.cell(cnt,4).text =ot
        if '输出延时' in str1:
            od=str1[str1.rfind('：')+1:-2]
            table.cell(cnt,7).text =od
        if '短路无输出' in str1:
            if '正常' in str1:
                table.cell(cnt,6).text  ="正常"
            else:
                table.cell(cnt,6).text  =" 不正常"
        if '欠压闪灯10次' in str1:
            if '正常' in str1:
                table2.cell(5,1).text ="闪灯10下"
            else:
                table2.cell(5,2).text =table2.cell(5,2).text+ str(cnt) +" 闪灯异常"
        if '上电闪烁次数' in str1:
            if '1' in str1:
                table2.cell(1,1).text ="闪灯1下"
            elif '3' in str1:
                table2.cell(1,2).text =table2.cell(1,2).text+ str(cnt) +"闪灯3下 "
        if '超时保护闪灯两次' in str1:
            if '正常' in str1:
                table2.cell(3,1).text ="闪2下"
            else:
                table2.cell(3,2).text =table2.cell(3,2).text+ str(cnt) +" 闪灯异常 "
        if '短路长亮2S' in str1:
            if '正常' in str1:
                table2.cell(4,1).text ="长亮2S"
            else:
                table2.cell(4,2).text =table2.cell(4,2).text+ str(cnt) +" 闪灯异常 "
        if '吸烟渐亮' in str1:
            if '正常' in str1:
                gradual_high=1
            else:
                gradual_high=0
        if '停止吸烟渐灭' in str1:
            if '正常' in str1:
                gradual_low=1
            else:
                gradual_low=0
        if '完成' in str1:
            if gradual_high==1 and gradual_low==1:
                table2.cell(2,1).text ="渐亮渐灭"
            else:
                table2.cell(2,2).text =table2.cell(2,2).text+ str(cnt) + + "灯异常 "
            #doc.save('E:/E/PM/测试报告/test result report.docx')
            doc.save('test result report.docx')
            break;
        if 'flash ok' in str1 and i>6:
            cnt=cnt-1;
            break;
            
com.close()

#doc.save('E:/E/PM/测试报告/test result report.docx')
print('测试完成')